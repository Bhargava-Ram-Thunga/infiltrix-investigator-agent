"""RAG engine for the Investigator Agent.

Two independent enrichment capabilities:

1. **Handbook snippet retrieval** — ingests the Infiltrix AI Engineering
   Handbook (DOCX) once at construction and returns the most relevant
   paragraph for a threat type via keyword scoring.
2. **MITRE ATT&CK mapping** — a static, deterministic table from IDS threat
   identifiers to technique codes. This works even when the handbook file
   is missing; the two features degrade independently.

Retrieval approach (and why it's this simple)
---------------------------------------------
The handbook is a general AI-engineering text, not a threat-intel corpus,
so heavyweight retrieval (embeddings, vector stores) buys nothing here.
Scoring = sum of keyword occurrence counts per paragraph; highest score
wins; ties keep the earliest paragraph. This is O(paragraphs × keywords)
over ~260 paragraphs — microseconds, no index needed.

Upgrading retrieval later
-------------------------
Keep the public surface (`search`, `map_to_mitre`, `enrich`) stable and the
agent needs no changes. To move to embeddings: replace `_ingest`/`search`
internals, precompute vectors in `__init__`, and keep returning
`Optional[str]` capped at `max_chars`.

Adding a new threat type
------------------------
1. Add the identifier to `MITRE_MAPPING` with its (code, name) tuple.
2. Add a keyword list to `_THREAT_KEYWORDS` (lowercase; scored by substring
   count, so prefer distinctive stems like "scan" over generic words).
3. Add the parametrized cases in tests/test_rag.py — the
   `test_mapping_covers_all_spec_threats` test will remind you.
No schema or agent changes are required: threat identifiers flow through as
plain strings.
"""

from pathlib import Path
from typing import Dict, List, Optional, Tuple

try:
    from docx import Document
except ImportError:  # pragma: no cover - python-docx is a hard requirement
    Document = None

_MODULE_DIR = Path(__file__).resolve().parent
_HANDBOOK_NAME = "Artificial_Intelligence_Engineering_Handbook_INFILTRIX.docx"

# Handbook discovery order (first hit wins):
#   1. ./docs/              — primary; the module ships with its copy here
#   2. <parent root>/docs/  — shared project copy (fallback)
#   3. <parent root>/       — legacy flat layout (fallback)
# An explicit path passed to RAGEngine(handbook_path=...) bypasses this list.
HANDBOOK_CANDIDATE_PATHS = [
    _MODULE_DIR / "docs" / _HANDBOOK_NAME,
    _MODULE_DIR.parent / "docs" / _HANDBOOK_NAME,
    _MODULE_DIR.parent / _HANDBOOK_NAME,
]

# Threat identifier -> (MITRE ATT&CK code, human-readable technique name).
# This is the authoritative enrichment; handbook snippets are best-effort
# context. Codes are per the INFILTRIX engineering specification.
MITRE_MAPPING: Dict[str, Tuple[str, str]] = {
    "credential_stuffing": ("T1110.004", "Credential Stuffing"),
    "endpoint_scan": ("T1595.002", "Vulnerability Scanning"),
    "high_rate_api_abuse": ("T1499.002", "Service Exhaustion Flood"),
    "status_code_anomaly": ("T1083", "File and Directory Discovery"),
}

# Keywords used to score handbook paragraphs per threat type. All lowercase;
# matching is case-insensitive substring counting. Tune these (not the
# scoring algorithm) first when a threat retrieves an irrelevant snippet.
_THREAT_KEYWORDS: Dict[str, List[str]] = {
    "credential_stuffing": ["credential", "authentication", "login", "brute", "password", "attack"],
    "endpoint_scan": ["scan", "reconnaissance", "vulnerability", "probe", "network", "traffic"],
    "high_rate_api_abuse": ["ddos", "flood", "rate", "abuse", "exhaustion", "denial", "traffic"],
    "status_code_anomaly": ["anomaly", "anomalies", "detection", "reconstruction", "unusual", "behavior"],
}

# Paragraphs shorter than this are skipped at ingestion — filters out
# headings, list markers, and boilerplate that would otherwise win on a
# single keyword hit.
_MIN_PARAGRAPH_CHARS = 40


class RAGEngine:
    """Handbook snippet retrieval and MITRE ATT&CK threat mapping.

    Construction cost: one full DOCX parse (~hundreds of ms). Build ONE
    instance at service startup and share it — it is immutable after
    construction and therefore thread-safe to read from concurrently.

    Degradation contract (never raises for a missing handbook):
    * handbook missing/unreadable  -> `search()` returns None, mapping works
    * unknown threat identifier    -> both `search()` and `map_to_mitre()`
                                      return None
    """

    def __init__(self, handbook_path: Optional[str] = None):
        """Ingest the handbook.

        Args:
            handbook_path: Explicit DOCX path. When None, the candidate
                paths in HANDBOOK_CANDIDATE_PATHS are probed in order. A
                non-existent explicit path resolves to None (degraded mode)
                rather than raising — deployments without the handbook must
                still boot.
        """
        self.handbook_path = self._resolve_handbook_path(handbook_path)
        self.paragraphs: List[str] = self._ingest()

    @staticmethod
    def _resolve_handbook_path(handbook_path: Optional[str]) -> Optional[Path]:
        if handbook_path is not None:
            path = Path(handbook_path)
            return path if path.exists() else None
        for candidate in HANDBOOK_CANDIDATE_PATHS:
            if candidate.exists():
                return candidate
        return None

    def _ingest(self) -> List[str]:
        """Parse the DOCX into substantial paragraphs (see _MIN_PARAGRAPH_CHARS)."""
        if self.handbook_path is None or Document is None:
            return []
        document = Document(str(self.handbook_path))
        return [
            p.text.strip()
            for p in document.paragraphs
            if len(p.text.strip()) >= _MIN_PARAGRAPH_CHARS
        ]

    def map_to_mitre(self, threat: str) -> Optional[str]:
        """Map an IDS threat identifier to its MITRE ATT&CK technique code.

        Returns None for unregistered threats — callers must handle None
        rather than assuming every detection maps to a technique.
        """
        mapping = MITRE_MAPPING.get(threat)
        return mapping[0] if mapping else None

    def search(self, threat: str, max_chars: int = 500) -> Optional[str]:
        """Return the highest-scoring handbook paragraph for the threat type.

        Args:
            threat: IDS threat identifier; must have a `_THREAT_KEYWORDS`
                entry to be searchable.
            max_chars: Hard cap on the returned snippet — keeps
                InvestigationResult payloads bounded for downstream agents.

        Returns:
            The best paragraph truncated to `max_chars`, or None when the
            handbook is unavailable, the threat has no keywords, or no
            paragraph scored above zero.
        """
        keywords = _THREAT_KEYWORDS.get(threat)
        if not keywords or not self.paragraphs:
            return None
        best_score, best_paragraph = 0, None
        for paragraph in self.paragraphs:
            lowered = paragraph.lower()
            score = sum(lowered.count(keyword) for keyword in keywords)
            if score > best_score:
                best_score, best_paragraph = score, paragraph
        if best_paragraph is None:
            return None
        return best_paragraph[:max_chars]

    def enrich(self, threat: str) -> Tuple[Optional[str], Optional[str]]:
        """One-call enrichment: (handbook_snippet, mitre_technique).

        This is the method the agent calls; keep its signature stable when
        upgrading retrieval internals.
        """
        return self.search(threat), self.map_to_mitre(threat)
